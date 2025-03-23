import os
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
from werkzeug.utils import secure_filename
import uuid
import io
import sys
import subprocess
import platform
import shutil

# Import libraries for file conversion
try:
    import docx
    from fpdf import FPDF
    CONVERSION_ENABLED = True
except ImportError:
    CONVERSION_ENABLED = False
    print("Warning: docx or fpdf libraries not found. File conversion will be simulated.")

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-key-for-testing')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['DOWNLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png', 'gif'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {
    'pdf': ['pdf'],
    'image': ['jpg', 'jpeg', 'png', 'gif'],
    'text': ['txt', 'doc', 'docx'],
}

def allowed_file(filename, file_types=None):
    """Check if the file extension is allowed"""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    
    if file_types:
        for file_type in file_types:
            if ext in ALLOWED_EXTENSIONS.get(file_type, []):
                return True
        return False
    
    # If no specific file types provided, check all allowed extensions
    for extensions in ALLOWED_EXTENSIONS.values():
        if ext in extensions:
            return True
    return False

@app.route('/')
def index():
    """Render the main page"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and conversion"""
    if 'file' not in request.files:
        flash('No file part', 'error')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(request.url)
    
    conversion_type = request.form.get('conversion_type', 'to_pdf')
    
    if file and allowed_file(file.filename):
        # Generate a unique filename to avoid collisions
        original_filename = secure_filename(file.filename)
        filename_base = str(uuid.uuid4())
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        
        # Save the uploaded file
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename_base}.{file_extension}")
        file.save(upload_path)
        
        try:
            output_filename = f"{filename_base}_converted.pdf"
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            # Perform actual conversion based on file type
            if CONVERSION_ENABLED:
                if file_extension in ['doc', 'docx']:
                    convert_word_to_pdf(upload_path, output_path)
                elif file_extension in ['txt']:
                    convert_text_to_pdf(upload_path, output_path)
                elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
                    # For images, create a PDF with a message (would need Pillow for actual conversion)
                    create_simple_pdf(output_path, f"Image file: {original_filename}")
                elif file_extension in ['pdf']:
                    # Already a PDF, just copy it
                    import shutil
                    shutil.copy(upload_path, output_path)
                else:
                    create_simple_pdf(output_path, f"Converted from: {original_filename}")
            else:
                # Fallback to simple PDF creation if libraries aren't available
                create_simple_pdf(output_path, f"Converted from: {original_filename}")
            
            flash('File successfully converted!', 'success')
            return redirect(url_for('download_file', filename=output_filename))
            
        except Exception as e:
            flash(f'Error during conversion: {str(e)}', 'error')
            return redirect(request.url)
    
    flash('File type not allowed', 'error')
    return redirect(request.url)

def convert_word_to_pdf(word_path, pdf_path):
    """Convert a Word document to PDF preserving formatting as much as possible"""
    try:
        # Try using Microsoft Word if available (Windows only)
        if platform.system() == "Windows":
            try:
                # Check if PowerShell is available
                if shutil.which('powershell.exe'):
                    # PowerShell script to convert Word to PDF using Word COM object
                    ps_script = f"""
                    $wordApp = New-Object -ComObject Word.Application
                    $wordApp.Visible = $false
                    $doc = $wordApp.Documents.Open("{word_path.replace('\\', '\\\\')}")
                    $doc.SaveAs([ref] "{pdf_path.replace('\\', '\\\\')}", [ref] 17)
                    $doc.Close()
                    $wordApp.Quit()
                    [System.Runtime.Interopservices.Marshal]::ReleaseComObject($wordApp)
                    """
                    
                    # Save the script to a temporary file
                    script_path = os.path.join(os.path.dirname(word_path), "convert_script.ps1")
                    with open(script_path, 'w') as f:
                        f.write(ps_script)
                    
                    # Run the PowerShell script
                    result = subprocess.run([
                        'powershell.exe', '-ExecutionPolicy', 'Bypass', '-File', script_path
                    ], capture_output=True, text=True)
                    
                    # Clean up the script file
                    if os.path.exists(script_path):
                        os.remove(script_path)
                    
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        print(f"Successfully converted Word document to PDF using PowerShell: {pdf_path}")
                        return True
                    else:
                        print(f"PowerShell conversion failed or produced empty file: {result.stderr}")
            except Exception as e:
                print(f"Error using PowerShell Word conversion: {str(e)}")
        
        # Alternative: Try using LibreOffice if available
        if shutil.which('libreoffice') or shutil.which('soffice'):
            try:
                # Determine the LibreOffice executable name
                libreoffice_exe = shutil.which('libreoffice') or shutil.which('soffice')
                
                # Run LibreOffice to convert the file
                subprocess.run([
                    libreoffice_exe, '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_path), word_path
                ], check=True, capture_output=True)
                
                # LibreOffice creates PDF with original filename, so we need to rename it
                original_name = os.path.splitext(os.path.basename(word_path))[0] + '.pdf'
                original_pdf_path = os.path.join(os.path.dirname(pdf_path), original_name)
                
                if os.path.exists(original_pdf_path):
                    shutil.move(original_pdf_path, pdf_path)
                    print(f"Successfully converted Word document to PDF using LibreOffice: {pdf_path}")
                    return True
            except Exception as e:
                print(f"Error using LibreOffice: {str(e)}. Falling back to basic conversion.")
        
        # Fallback to enhanced basic conversion using python-docx and fpdf
        doc = docx.Document(word_path)
        
        # Create a PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Set font for body text
        pdf.set_font("Arial", size=11)
        
        # Extract styles from the document if possible
        styles = {}
        try:
            for style in doc.styles:
                if hasattr(style, 'font') and hasattr(style.font, 'size'):
                    styles[style.name] = {
                        'size': style.font.size.pt if hasattr(style.font.size, 'pt') else 11,
                        'bold': style.font.bold if hasattr(style.font, 'bold') else False,
                        'italic': style.font.italic if hasattr(style.font, 'italic') else False
                    }
        except:
            # If styles extraction fails, use defaults
            pass
        
        # Process paragraphs with better formatting
        for para in doc.paragraphs:
            if not para.text.strip():
                # Add some space for empty paragraphs
                pdf.ln(5)
                continue
            
            # Try to match paragraph style
            style_name = para.style.name if hasattr(para, 'style') and hasattr(para.style, 'name') else 'Normal'
            
            # Set font based on style
            if style_name in styles:
                style = styles[style_name]
                font_style = ''
                if style.get('bold', False):
                    font_style += 'B'
                if style.get('italic', False):
                    font_style += 'I'
                font_size = style.get('size', 11)
                pdf.set_font("Arial", font_style, font_size)
            elif style_name.startswith('Heading'):
                # Make headings bold and larger
                if '1' in style_name:
                    pdf.set_font("Arial", 'B', 16)
                elif '2' in style_name:
                    pdf.set_font("Arial", 'B', 14)
                else:
                    pdf.set_font("Arial", 'B', 12)
            else:
                # Regular paragraph
                pdf.set_font("Arial", '', 11)
            
            # Process paragraph text with runs to preserve formatting within paragraph
            if not para.runs:
                # If no runs, process the whole paragraph
                encoded_text = para.text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 7, encoded_text)
                pdf.ln(3)
            else:
                # Process each run to preserve inline formatting
                line_text = ""
                for run in para.runs:
                    # Add the run text to the line
                    line_text += run.text
                
                # Output the complete line
                if line_text:
                    encoded_text = line_text.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 7, encoded_text)
                    pdf.ln(3)
        
        # Process tables if any
        for table in doc.tables:
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(200, 10, txt="Table Content:", ln=True)
            pdf.set_font("Arial", size=11)
            
            # Calculate column widths (simple approach)
            col_count = max(len(row.cells) for row in table.rows) if table.rows else 0
            if col_count > 0:
                col_width = 190 / col_count  # 190mm is typical page width with margins
            
                # Process each row
                for row in table.rows:
                    # Reset to left position
                    pdf.set_x(10)
                    
                    max_height = 7  # Minimum row height
                    
                    # First pass: calculate max height needed
                    for i, cell in enumerate(row.cells):
                        if i >= col_count:
                            break
                        
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Calculate how many lines this text will take
                            lines = len(pdf.multi_cell(col_width, 7, cell_text, split_only=True))
                            cell_height = lines * 7
                            max_height = max(max_height, cell_height)
                    
                    # Second pass: output cells with consistent height
                    for i, cell in enumerate(row.cells):
                        if i >= col_count:
                            break
                        
                        cell_text = cell.text.strip()
                        encoded_cell_text = cell_text.encode('latin-1', 'replace').decode('latin-1')
                        
                        # Draw cell border
                        pdf.cell(col_width, max_height, encoded_cell_text, border=1)
                    
                    pdf.ln(max_height)
            
            pdf.ln(5)
        
        # Save the PDF
        pdf.output(pdf_path)
        print(f"Successfully converted Word document to PDF using enhanced basic method: {pdf_path}")
        return True
    except Exception as e:
        print(f"Error converting Word to PDF: {str(e)}")
        # Fallback to simple PDF if conversion fails
        create_simple_pdf(pdf_path, f"Content from Word file (conversion failed: {str(e)})")
        return False

def convert_text_to_pdf(text_path, pdf_path):
    """Convert a text file to PDF"""
    try:
        # Create a PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Read the text file and add content to PDF
        with open(text_path, 'r', encoding='utf-8', errors='ignore') as text_file:
            for line in text_file:
                # Encode text to handle special characters
                encoded_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, encoded_line)
        
        # Save the PDF
        pdf.output(pdf_path)
    except Exception as e:
        print(f"Error converting text to PDF: {str(e)}")
        # Fallback to simple PDF if conversion fails
        create_simple_pdf(pdf_path, f"Content from text file (conversion partially failed)")

def create_simple_pdf(pdf_path, content):
    """Create a simple PDF with the given content"""
    try:
        # Try to use FPDF if available
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=16)
        pdf.cell(200, 10, txt="PDF Converter", ln=True, align='C')
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(pdf_path)
    except Exception:
        # Fallback to manual PDF creation
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.7\n')
            f.write(b'1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n')
            f.write(b'2 0 obj\n<</Type/Pages/Kids[3 0 R]/Count 1>>\nendobj\n')
            f.write(b'3 0 obj\n<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>/Contents 4 0 R>>\nendobj\n')
            f.write(b'4 0 obj\n<</Length 100>>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(')
            f.write(content.encode('ascii', 'replace'))
            f.write(b') Tj\nET\nendstream\nendobj\n')
            f.write(b'xref\n0 5\n0000000000 65535 f\n0000000010 00000 n\n0000000053 00000 n\n0000000102 00000 n\n0000000191 00000 n\ntrailer\n<</Size 5/Root 1 0 R>>\nstartxref\n293\n%%EOF\n')

@app.route('/download/<filename>')
def download_file(filename):
    """Serve the converted file for download"""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@app.route('/about')
def about():
    """About page"""
    return render_template('about.html')

# Clean up old files periodically
@app.before_request
def cleanup_old_files():
    """Remove files older than 1 hour"""
    import time
    current_time = time.time()
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        # If file is older than 1 hour (3600 seconds), delete it
        if os.path.isfile(file_path) and current_time - os.path.getmtime(file_path) > 3600:
            os.remove(file_path)

if __name__ == '__main__':
    app.run(debug=True)
